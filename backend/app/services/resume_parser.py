"""
Resume parsing functionality for portfolio analysis.

Implements Requirements:
- 13.5: Resume file parsing (PDF, DOCX, TXT formats)
- 13.6: NLP skill extraction with spaCy
"""
import io
import re
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import PyPDF2
import docx
import spacy

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parser for extracting information from resume files."""
    
    def __init__(self):
        """Initialize resume parser with spaCy NLP model."""
        try:
            # Load spaCy English model for NLP
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model 'en_core_web_sm' not found. Run: python -m spacy download en_core_web_sm")
            self.nlp = None
    
    def parse_resume(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """
        Parse resume file and extract structured information.
        
        Supports PDF, DOCX, and TXT formats. Extracts skills, experience,
        education, and other relevant information using NLP.
        
        Implements Requirements:
        - 13.5: Parse PDF, DOCX, TXT formats
        - 13.6: Use NLP for skill extraction
        
        Args:
            file_content: Raw file content as bytes
            file_type: File type ('pdf', 'docx', 'txt')
            
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - skills: List of detected technical skills
                - experience: List of work experience entries
                - education: List of education entries
                - contact_info: Extracted contact information
                - proficiency_levels: Dict of skill: proficiency
                - experience_years: Estimated years of experience
                
        Raises:
            ValueError: If file type is unsupported or parsing fails
        """
        # Extract text based on file type
        file_type = file_type.lower().strip('.')
        
        if file_type == 'pdf':
            text = self._extract_text_from_pdf(file_content)
        elif file_type in ['docx', 'doc']:
            text = self._extract_text_from_docx(file_content)
        elif file_type == 'txt':
            text = self._extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, TXT")
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Resume file appears to be empty or too short")
        
        # Extract structured information
        contact_info = self._extract_contact_info(text)
        skills = self._extract_skills_nlp(text)
        experience = self._extract_experience(text)
        education = self._extract_education(text)
        
        # Calculate proficiency levels based on context
        proficiency_levels = self._calculate_skill_proficiency(text, skills)
        
        # Estimate years of experience
        experience_years = self._estimate_experience_years(experience, text)
        
        return {
            "text": text,
            "skills": skills,
            "experience": experience,
            "education": education,
            "contact_info": contact_info,
            "proficiency_levels": proficiency_levels,
            "experience_years": experience_years
        }
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If PDF parsing fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            text = "\n".join(text_parts)
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to parse PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If DOCX parsing fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    def _extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_content: TXT file content as bytes
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If TXT parsing fails
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to parse TXT: {str(e)}")
            raise ValueError(f"Failed to parse TXT file: {str(e)}")

    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """
        Extract contact information from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with email, phone, linkedin, github
        """
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None
        }
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group(0)
        
        # Extract phone (various formats)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group(0)
        
        # Extract LinkedIn URL
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info["linkedin"] = f"https://{linkedin_match.group(0)}"
        
        # Extract GitHub URL
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info["github"] = f"https://{github_match.group(0)}"
        
        return contact_info

    
    def _extract_skills_nlp(self, text: str) -> List[str]:
        """
        Extract technical skills using NLP and pattern matching.
        
        Implements Requirement 13.6: Use NLP for skill extraction.
        
        Args:
            text: Resume text
            
        Returns:
            List of detected technical skills
        """
        skills = set()
        
        # Comprehensive list of technical skills to detect
        skill_keywords = {
            # Programming Languages
            "python", "java", "javascript", "typescript", "c++", "c#", "c", "ruby", 
            "go", "golang", "rust", "php", "swift", "kotlin", "scala", "r", "matlab",
            "perl", "shell", "bash", "powershell", "objective-c", "dart", "elixir",
            "haskell", "clojure", "groovy", "lua", "vb.net", "f#",
            
            # Web Frameworks & Libraries
            "react", "reactjs", "react.js", "angular", "angularjs", "vue", "vuejs", 
            "vue.js", "node", "nodejs", "node.js", "express", "expressjs", "django",
            "flask", "fastapi", "spring", "spring boot", "asp.net", ".net", "dotnet",
            "laravel", "symfony", "rails", "ruby on rails", "nextjs", "next.js",
            "nuxt", "svelte", "ember", "backbone", "jquery",
            
            # Mobile Development
            "android", "ios", "react native", "flutter", "xamarin", "ionic",
            "cordova", "phonegap", "swiftui",
            
            # Cloud Platforms
            "aws", "amazon web services", "azure", "microsoft azure", "gcp", 
            "google cloud", "google cloud platform", "heroku", "digitalocean",
            "linode", "cloudflare", "vercel", "netlify",
            
            # DevOps & Tools
            "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins",
            "gitlab", "github actions", "circleci", "travis ci", "bamboo",
            "puppet", "chef", "vagrant", "helm", "istio", "prometheus",
            "grafana", "elk", "elasticsearch", "logstash", "kibana",
            
            # Databases
            "sql", "postgresql", "postgres", "mysql", "mongodb", "redis",
            "cassandra", "dynamodb", "oracle", "sql server", "mariadb",
            "sqlite", "couchdb", "neo4j", "influxdb", "timescaledb",
            "firestore", "cosmos db",
            
            # Data Science & ML
            "machine learning", "deep learning", "ai", "artificial intelligence",
            "data science", "nlp", "natural language processing", "computer vision",
            "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas",
            "numpy", "scipy", "matplotlib", "seaborn", "jupyter", "spark",
            "hadoop", "kafka", "airflow", "mlflow", "kubeflow",
            
            # Testing
            "pytest", "unittest", "jest", "mocha", "jasmine", "selenium",
            "cypress", "junit", "testng", "rspec", "cucumber", "postman",
            
            # API & Architecture
            "rest", "restful", "graphql", "grpc", "soap", "api", "microservices",
            "serverless", "lambda", "event-driven", "message queue", "rabbitmq",
            "sqs", "sns", "pub/sub",
            
            # Version Control
            "git", "github", "gitlab", "bitbucket", "svn", "mercurial",
            
            # Methodologies
            "agile", "scrum", "kanban", "devops", "ci/cd", "tdd", "bdd",
            "pair programming", "code review",
            
            # Other Technologies
            "html", "css", "sass", "scss", "less", "webpack", "babel",
            "typescript", "graphql", "redux", "mobx", "rxjs", "websocket",
            "oauth", "jwt", "saml", "ldap", "active directory"
        }
        
        # Convert text to lowercase for matching
        text_lower = text.lower()
        
        # Pattern matching for skills
        for skill in skill_keywords:
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                # Normalize skill name
                normalized_skill = self._normalize_skill_name(skill)
                skills.add(normalized_skill)
        
        # Use spaCy NLP for additional entity extraction if available
        if self.nlp:
            try:
                doc = self.nlp(text[:10000])  # Limit text length for performance
                
                # Extract noun chunks that might be skills
                for chunk in doc.noun_chunks:
                    chunk_text = chunk.text.lower().strip()
                    # Check if chunk matches known skills
                    if chunk_text in skill_keywords:
                        normalized_skill = self._normalize_skill_name(chunk_text)
                        skills.add(normalized_skill)
                
            except Exception as e:
                logger.warning(f"spaCy NLP extraction failed: {str(e)}")
        
        return sorted(list(skills))

    
    def _normalize_skill_name(self, skill: str) -> str:
        """
        Normalize skill name for consistency.
        
        Args:
            skill: Raw skill name
            
        Returns:
            Normalized skill name
        """
        # Special cases for common variations
        normalizations = {
            "reactjs": "React",
            "react.js": "React",
            "angularjs": "Angular",
            "vuejs": "Vue.js",
            "vue.js": "Vue.js",
            "nodejs": "Node.js",
            "node.js": "Node.js",
            "expressjs": "Express",
            "spring boot": "Spring Boot",
            "asp.net": "ASP.NET",
            ".net": ".NET",
            "dotnet": ".NET",
            "ruby on rails": "Ruby on Rails",
            "nextjs": "Next.js",
            "next.js": "Next.js",
            "react native": "React Native",
            "amazon web services": "AWS",
            "microsoft azure": "Azure",
            "google cloud platform": "GCP",
            "google cloud": "GCP",
            "k8s": "Kubernetes",
            "postgres": "PostgreSQL",
            "sql server": "SQL Server",
            "sklearn": "Scikit-learn",
            "scikit-learn": "Scikit-learn",
            "machine learning": "Machine Learning",
            "deep learning": "Deep Learning",
            "artificial intelligence": "AI",
            "natural language processing": "NLP",
            "computer vision": "Computer Vision",
            "ci/cd": "CI/CD",
            "tdd": "TDD",
            "bdd": "BDD"
        }
        
        skill_lower = skill.lower().strip()
        if skill_lower in normalizations:
            return normalizations[skill_lower]
        
        # Default: title case
        return skill.title()

    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract work experience entries from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of experience dictionaries with title, company, dates, description
        """
        experiences = []
        
        # Common section headers for experience
        experience_headers = [
            r'work\s+experience',
            r'professional\s+experience',
            r'employment\s+history',
            r'experience',
            r'work\s+history',
            r'career\s+history'
        ]
        
        # Find experience section
        experience_section = None
        text_lower = text.lower()
        
        for header in experience_headers:
            match = re.search(header, text_lower)
            if match:
                # Extract text after this header
                start_pos = match.end()
                # Find next major section (education, skills, etc.)
                next_section_pattern = r'\n\s*(education|skills|certifications|projects|awards)'
                next_match = re.search(next_section_pattern, text_lower[start_pos:])
                
                if next_match:
                    end_pos = start_pos + next_match.start()
                    experience_section = text[start_pos:end_pos]
                else:
                    experience_section = text[start_pos:]
                break
        
        if not experience_section:
            # Try to find experience entries without explicit section
            experience_section = text
        
        # Pattern for dates (various formats)
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)',
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*[-–—]\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}',
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*[-–—]\s*(present|current)',
            r'\d{1,2}/\d{4}\s*[-–—]\s*\d{1,2}/\d{4}',
            r'\d{1,2}/\d{4}\s*[-–—]\s*(present|current)'
        ]
        
        # Split into potential experience entries (by date patterns)
        lines = experience_section.split('\n')
        current_entry = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains a date pattern
            has_date = False
            for pattern in date_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    has_date = True
                    # Save previous entry if exists
                    if current_entry and current_entry.get("title"):
                        experiences.append(current_entry)
                    
                    # Start new entry
                    current_entry = {
                        "title": "",
                        "company": "",
                        "dates": line,
                        "description": ""
                    }
                    break
            
            # If we have a current entry, add content to it
            if current_entry:
                if not current_entry["title"] and not has_date:
                    # First line after dates is likely the title
                    current_entry["title"] = line
                elif not current_entry["company"] and current_entry["title"] and not has_date:
                    # Second line is likely the company
                    current_entry["company"] = line
                elif not has_date:
                    # Subsequent lines are description
                    if current_entry["description"]:
                        current_entry["description"] += " " + line
                    else:
                        current_entry["description"] = line
        
        # Add last entry
        if current_entry and current_entry.get("title"):
            experiences.append(current_entry)
        
        return experiences

    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract education entries from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            List of education dictionaries with degree, school, year
        """
        education_entries = []
        
        # Common section headers for education
        education_headers = [
            r'education',
            r'academic\s+background',
            r'educational\s+background',
            r'qualifications'
        ]
        
        # Find education section
        education_section = None
        text_lower = text.lower()
        
        for header in education_headers:
            match = re.search(header, text_lower)
            if match:
                start_pos = match.end()
                # Find next major section
                next_section_pattern = r'\n\s*(experience|work|skills|certifications|projects|awards)'
                next_match = re.search(next_section_pattern, text_lower[start_pos:])
                
                if next_match:
                    end_pos = start_pos + next_match.start()
                    education_section = text[start_pos:end_pos]
                else:
                    education_section = text[start_pos:]
                break
        
        if not education_section:
            return education_entries
        
        # Common degree keywords
        degree_keywords = [
            r'ph\.?d', r'doctorate', r'doctor of philosophy',
            r'master', r'm\.?s\.?', r'm\.?a\.?', r'mba', r'm\.?eng',
            r'bachelor', r'b\.?s\.?', r'b\.?a\.?', r'b\.?eng', r'b\.?tech',
            r'associate', r'a\.?s\.?', r'a\.?a\.?',
            r'diploma', r'certificate'
        ]
        
        # Extract education entries
        lines = education_section.split('\n')
        current_entry = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line contains a degree keyword
            has_degree = False
            for degree_pattern in degree_keywords:
                if re.search(degree_pattern, line, re.IGNORECASE):
                    has_degree = True
                    # Save previous entry
                    if current_entry and current_entry.get("degree"):
                        education_entries.append(current_entry)
                    
                    # Start new entry
                    current_entry = {
                        "degree": line,
                        "school": "",
                        "year": ""
                    }
                    
                    # Try to extract year from same line
                    year_match = re.search(r'\b(19|20)\d{2}\b', line)
                    if year_match:
                        current_entry["year"] = year_match.group(0)
                    
                    break
            
            # If we have a current entry and this line doesn't have a degree
            if current_entry and not has_degree:
                if not current_entry["school"]:
                    # Next line after degree is likely the school
                    current_entry["school"] = line
                    # Try to extract year if not already found
                    if not current_entry["year"]:
                        year_match = re.search(r'\b(19|20)\d{2}\b', line)
                        if year_match:
                            current_entry["year"] = year_match.group(0)
        
        # Add last entry
        if current_entry and current_entry.get("degree"):
            education_entries.append(current_entry)
        
        return education_entries

    
    def _calculate_skill_proficiency(self, text: str, skills: List[str]) -> Dict[str, float]:
        """
        Calculate proficiency levels for detected skills based on context.
        
        Implements Requirement 13.6: Identify proficiency levels using NLP.
        
        Args:
            text: Resume text
            skills: List of detected skills
            
        Returns:
            Dictionary mapping skill to proficiency level (0.0-1.0)
        """
        proficiency_levels = {}
        text_lower = text.lower()
        
        # Proficiency keywords and their weights
        proficiency_keywords = {
            "expert": 1.0,
            "advanced": 0.9,
            "proficient": 0.8,
            "experienced": 0.8,
            "senior": 0.85,
            "lead": 0.9,
            "architect": 0.95,
            "intermediate": 0.6,
            "competent": 0.6,
            "familiar": 0.4,
            "basic": 0.3,
            "beginner": 0.2,
            "learning": 0.2
        }
        
        for skill in skills:
            skill_lower = skill.lower()
            proficiency = 0.5  # Default medium proficiency
            
            # Find skill mentions in text
            skill_pattern = r'\b' + re.escape(skill_lower) + r'\b'
            matches = list(re.finditer(skill_pattern, text_lower))
            
            if matches:
                # Check context around each mention
                max_proficiency = 0.5
                
                for match in matches:
                    # Get context (50 characters before and after)
                    start = max(0, match.start() - 50)
                    end = min(len(text_lower), match.end() + 50)
                    context = text_lower[start:end]
                    
                    # Check for proficiency keywords in context
                    for keyword, weight in proficiency_keywords.items():
                        if keyword in context:
                            max_proficiency = max(max_proficiency, weight)
                
                # Count mentions (more mentions = higher proficiency)
                mention_count = len(matches)
                mention_bonus = min(0.2, mention_count * 0.05)
                
                proficiency = min(1.0, max_proficiency + mention_bonus)
            
            proficiency_levels[skill] = round(proficiency, 2)
        
        return proficiency_levels
    
    def _estimate_experience_years(self, experience: List[Dict[str, Any]], text: str) -> float:
        """
        Estimate total years of professional experience.
        
        Args:
            experience: List of experience entries
            text: Resume text
            
        Returns:
            Estimated years of experience
        """
        if not experience:
            # Try to find years mentioned in text
            years_pattern = r'(\d+)\+?\s*years?\s+(?:of\s+)?experience'
            match = re.search(years_pattern, text.lower())
            if match:
                return float(match.group(1))
            return 0.0
        
        total_years = 0.0
        
        for exp in experience:
            dates_str = exp.get("dates", "")
            
            # Extract years from date string
            year_matches = re.findall(r'\b(19|20)\d{2}\b', dates_str)
            
            if len(year_matches) >= 2:
                # Calculate duration
                start_year = int(year_matches[0])
                end_year = int(year_matches[-1])
                duration = end_year - start_year
                total_years += max(0, duration)
            elif len(year_matches) == 1:
                # Check if it's current/present
                if re.search(r'present|current', dates_str.lower()):
                    start_year = int(year_matches[0])
                    current_year = datetime.now().year
                    duration = current_year - start_year
                    total_years += max(0, duration)
                else:
                    # Assume 1 year if only one year mentioned
                    total_years += 1.0
        
        # Cap at reasonable maximum
        return min(total_years, 50.0)
